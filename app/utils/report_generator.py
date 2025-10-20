"""
PCB清洁生产审核报告生成器
生成包含企业信息、筹划与组织、预审核、审核数据的Word报告
"""
import os
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from app.models.pcb import PCBEnterprise, PCBPreAuditData, PCBAuditResult, PCBAuditReport
from app.models.pcb_planning import PCBLeadershipTeam, PCBWorkTeam, PCBWorkPlan, PCBTrainingRecord


class PCBReportGenerator:
    """PCB清洁生产审核报告生成器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()
    
    def setup_document_style(self):
        """设置文档样式"""
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
    
    def add_title(self, text: str, level: int = 1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return heading
    
    def add_paragraph(self, text: str, style: str = 'Normal'):
        """添加段落"""
        paragraph = self.doc.add_paragraph(text, style=style)
        return paragraph
    
    def add_table(self, data: List[List], headers: List[str] = None, title: str = None):
        """添加表格"""
        if title:
            self.add_paragraph(title, 'Heading 3')
        
        if headers:
            table = self.doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 设置表头
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                # 设置表头样式
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
        else:
            table = self.doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 添加数据行
        if headers and data:
            for row_data in data:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data) if cell_data is not None else ''
        elif data:
            for i, row_data in enumerate(data):
                if i < len(table.rows):
                    row_cells = table.rows[i].cells
                    for j, cell_data in enumerate(row_data):
                        if j < len(row_cells):
                            row_cells[j].text = str(cell_data) if cell_data is not None else ''
        
        return table
    
    async def generate_report(self, enterprise_id: int) -> str:
        """生成完整的审核报告"""
        # 获取企业信息
        enterprise = await PCBEnterprise.get_or_none(id=enterprise_id)
        if not enterprise:
            raise ValueError(f"企业ID {enterprise_id} 不存在")
        
        # 获取筹划与组织数据
        leadership_team = await PCBLeadershipTeam.filter(enterprise_id=enterprise_id)
        work_team = await PCBWorkTeam.filter(enterprise_id=enterprise_id)
        work_plans = await PCBWorkPlan.filter(enterprise_id=enterprise_id).order_by('stage_order')
        training_records = await PCBTrainingRecord.filter(enterprise_id=enterprise_id)
        
        # 获取预审核数据
        pre_audit_data = await PCBPreAuditData.get_or_none(enterprise_id=enterprise_id)
        
        # 获取审核结果
        audit_results = await PCBAuditResult.filter(enterprise_id=enterprise_id)
        
        # 获取审核报告
        audit_report = await PCBAuditReport.get_or_none(enterprise_id=enterprise_id)
        
        # 生成报告内容
        self._generate_cover_page(enterprise)
        self._generate_enterprise_info(enterprise)
        self._generate_planning_organization(leadership_team, work_team, work_plans, training_records)
        self._generate_pre_audit_data(pre_audit_data)
        await self._generate_audit_results(audit_results, audit_report)
        
        # 保存报告
        report_filename = f"PCB清洁生产审核报告_{enterprise.name}_{datetime.now().strftime('%Y%m%d')}.docx"
        report_path = os.path.join("reports", report_filename)
        
        # 确保reports目录存在
        os.makedirs("reports", exist_ok=True)
        
        self.doc.save(report_path)
        return report_path
    
    def _generate_cover_page(self, enterprise: PCBEnterprise):
        """生成封面页"""
        # 标题
        self.add_title("PCB行业清洁生产审核报告", 0)
        
        # 空行
        self.add_paragraph("")
        self.add_paragraph("")
        
        # 企业信息
        self.add_paragraph(f"企业名称：{enterprise.name}")
        self.add_paragraph(f"统一社会信用代码：{enterprise.unified_social_credit_code or '未填写'}")
        self.add_paragraph(f"所属地区：{enterprise.region or '未填写'} {enterprise.district or '未填写'}")
        self.add_paragraph(f"详细地址：{enterprise.address or '未填写'}")
        self.add_paragraph(f"法人代表：{enterprise.legal_representative or '未填写'}")
        self.add_paragraph(f"联系人：{enterprise.contact_person or '未填写'}")
        self.add_paragraph(f"联系电话：{enterprise.contact_phone or '未填写'}")
        self.add_paragraph(f"年产能：{enterprise.capacity or 0}万m²")
        
        # 空行
        self.add_paragraph("")
        self.add_paragraph("")
        
        # 报告信息
        self.add_paragraph(f"报告生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
        self.add_paragraph(f"审核状态：{enterprise.audit_status}")
        
        # 分页
        self.doc.add_page_break()
    
    def _generate_enterprise_info(self, enterprise: PCBEnterprise):
        """生成企业基本信息章节"""
        self.add_title("一、企业基本信息", 1)
        
        # 基本信息表格
        basic_info_data = [
            ["企业名称", enterprise.name],
            ["统一社会信用代码", enterprise.unified_social_credit_code or "未填写"],
            ["所属地市", enterprise.region or "未填写"],
            ["所属区县", enterprise.district or "未填写"],
            ["详细地址", enterprise.address or "未填写"],
            ["法人代表", enterprise.legal_representative or "未填写"],
            ["联系人", enterprise.contact_person or "未填写"],
            ["联系电话", enterprise.contact_phone or "未填写"],
            ["联系邮箱", enterprise.contact_email or "未填写"],
            ["行业类型", enterprise.industry_type or "未填写"],
            ["注册资本", f"{enterprise.capital or 0}万元"],
            ["年产能", f"{enterprise.capacity or 0}万m²"],
            ["审核状态", enterprise.audit_status],
            ["创建时间", enterprise.created_at.strftime('%Y-%m-%d %H:%M:%S') if enterprise.created_at else "未填写"]
        ]
        
        self.add_table(basic_info_data, ["项目", "内容"], "1.1 基本信息")
    
    def _generate_planning_organization(self, leadership_team: List, work_team: List, 
                                      work_plans: List, training_records: List):
        """生成筹划与组织章节"""
        self.add_title("二、筹划与组织", 1)
        
        # 领导小组
        if leadership_team:
            self.add_title("2.1 清洁生产审核领导小组", 2)
            leadership_data = []
            for member in leadership_team:
                leadership_data.append([
                    member.name,
                    member.position or "未填写",
                    member.department or "未填写",
                    member.role,
                    member.responsibilities or "未填写",
                    member.phone or "未填写"
                ])
            
            self.add_table(leadership_data, 
                          ["姓名", "职位", "部门", "角色", "职责", "联系电话"], 
                          "领导小组成员")
        
        # 工作小组
        if work_team:
            self.add_title("2.2 清洁生产审核工作小组", 2)
            work_team_data = []
            for member in work_team:
                work_team_data.append([
                    member.name,
                    member.position or "未填写",
                    member.department or "未填写",
                    member.role,
                    member.responsibilities or "未填写",
                    member.phone or "未填写"
                ])
            
            self.add_table(work_team_data, 
                          ["姓名", "职位", "部门", "角色", "职责", "联系电话"], 
                          "工作小组成员")
        
        # 工作计划
        if work_plans:
            self.add_title("2.3 清洁生产审核工作计划", 2)
            work_plan_data = []
            for plan in work_plans:
                work_plan_data.append([
                    f"第{plan.stage_order}阶段",
                    plan.stage,
                    plan.work_content or "未填写",
                    plan.planned_start_date.strftime('%Y-%m-%d') if plan.planned_start_date else "未填写",
                    plan.planned_end_date.strftime('%Y-%m-%d') if plan.planned_end_date else "未填写",
                    plan.responsible_department or "未填写",
                    plan.actual_start_date.strftime('%Y-%m-%d') if plan.actual_start_date else "未开始",
                    plan.actual_end_date.strftime('%Y-%m-%d') if plan.actual_end_date else "未完成"
                ])
            
            self.add_table(work_plan_data, 
                          ["阶段", "阶段名称", "工作内容", "计划开始时间", "计划结束时间", 
                           "责任部门", "实际开始时间", "实际结束时间"], 
                          "工作计划表")
        
        # 培训记录
        if training_records:
            self.add_title("2.4 培训记录", 2)
            training_data = []
            for record in training_records:
                training_data.append([
                    record.title,
                    record.date.strftime('%Y-%m-%d') if record.date else "未填写",
                    f"{record.duration or 0}分钟",
                    f"{record.participants or 0}人",
                    record.content or "未填写",
                    record.instructor or "未填写",
                    record.location or "未填写"
                ])
            
            self.add_table(training_data, 
                          ["培训标题", "培训日期", "培训时长", "参与人数", "培训内容", "培训讲师", "培训地点"], 
                          "培训记录表")
    
    def _generate_pre_audit_data(self, pre_audit_data: Optional[PCBPreAuditData]):
        """生成预审核数据章节"""
        self.add_title("三、预审核数据", 1)
        
        if not pre_audit_data:
            self.add_paragraph("暂无预审核数据")
            return
        
        # 生产情况
        if pre_audit_data.production_info:
            self.add_title("3.1 企业总体生产情况", 2)
            
            production_info = pre_audit_data.production_info
            if isinstance(production_info, dict):
                # 产能信息
                capacity_data = [["年产能", f"{production_info.get('capacity', 0)}万m²"]]
                self.add_table(capacity_data, ["项目", "数值"], "产能信息")
                
                # 产量信息
                if 'output' in production_info:
                    output_data = []
                    output_info = production_info['output']
                    if isinstance(output_info, dict):
                        for year, products in output_info.items():
                            if isinstance(products, dict):
                                for product_type, amount in products.items():
                                    output_data.append([year, product_type, f"{amount}万m²"])
                    
                    if output_data:
                        self.add_table(output_data, ["年份", "产品类型", "产量"], "产量统计")
        
        # 原辅材料使用情况
        if pre_audit_data.raw_materials:
            self.add_title("3.2 原辅材料使用情况", 2)
            
            raw_materials = pre_audit_data.raw_materials
            if isinstance(raw_materials, list):
                material_data = []
                for material in raw_materials:
                    if isinstance(material, dict):
                        material_data.append([
                            material.get('year', '未填写'),
                            material.get('name', '未填写'),
                            material.get('unit', '未填写'),
                            material.get('process', '未填写'),
                            material.get('amount', 0),
                            material.get('state', '未填写'),
                            f"{material.get('voc', 0)}%"
                        ])
                
                if material_data:
                    self.add_table(material_data, 
                                  ["年份", "材料名称", "单位", "工序", "用量", "状态", "VOC含量"], 
                                  "原辅材料使用情况")
        
        # 工艺装备
        if pre_audit_data.process_equipment:
            self.add_title("3.3 主要工艺及装备使用", 2)
            
            process_equipment = pre_audit_data.process_equipment
            if isinstance(process_equipment, dict):
                equipment_data = []
                
                # 刚性板工艺
                if 'rigid' in process_equipment:
                    rigid_info = process_equipment['rigid']
                    if isinstance(rigid_info, dict):
                        for product_type, info in rigid_info.items():
                            if isinstance(info, dict):
                                equipment_data.append([
                                    f"刚性{product_type}",
                                    info.get('line', '未填写'),
                                    info.get('process', '未填写'),
                                    info.get('equipment', '未填写')
                                ])
                
                # 挠性板工艺
                if 'flexible' in process_equipment:
                    flexible_info = process_equipment['flexible']
                    if isinstance(flexible_info, dict):
                        for product_type, info in flexible_info.items():
                            if isinstance(info, dict):
                                equipment_data.append([
                                    f"挠性{product_type}",
                                    info.get('line', '未填写'),
                                    info.get('process', '未填写'),
                                    info.get('equipment', '未填写')
                                ])
                
                if equipment_data:
                    self.add_table(equipment_data, 
                                  ["产品类型", "生产线", "工艺流程", "主要设备"], 
                                  "工艺装备情况")
        
        # 资源能源消耗
        if pre_audit_data.resource_consumption:
            self.add_title("3.4 资源能源消耗", 2)
            
            resource_consumption = pre_audit_data.resource_consumption
            if isinstance(resource_consumption, dict):
                # 用水情况
                if 'water' in resource_consumption:
                    water_data = []
                    water_info = resource_consumption['water']
                    if isinstance(water_info, list):
                        for water in water_info:
                            if isinstance(water, dict):
                                water_data.append([
                                    water.get('year', '未填写'),
                                    water.get('type', '未填写'),
                                    f"{water.get('amount', 0)}m³",
                                    water.get('source', '未填写')
                                ])
                    
                    if water_data:
                        self.add_table(water_data, 
                                      ["年份", "用水类型", "用水量", "水源"], 
                                      "用水情况统计")
                
                # 用电情况
                if 'electricity' in resource_consumption:
                    electricity_data = []
                    electricity_info = resource_consumption['electricity']
                    if isinstance(electricity_info, list):
                        for electricity in electricity_info:
                            if isinstance(electricity, dict):
                                electricity_data.append([
                                    electricity.get('year', '未填写'),
                                    electricity.get('type', '未填写'),
                                    f"{electricity.get('amount', 0)}kWh",
                                    electricity.get('source', '未填写')
                                ])
                    
                    if electricity_data:
                        self.add_table(electricity_data, 
                                      ["年份", "用电类型", "用电量", "电源"], 
                                      "用电情况统计")
        
        # 污染防治
        if pre_audit_data.pollution_control:
            self.add_title("3.5 污染防治", 2)
            
            pollution_control = pre_audit_data.pollution_control
            if isinstance(pollution_control, dict):
                # 铜回收率
                if 'copperRecovery' in pollution_control:
                    copper_data = []
                    copper_info = pollution_control['copperRecovery']
                    if isinstance(copper_info, list):
                        for copper in copper_info:
                            if isinstance(copper, dict):
                                copper_data.append([
                                    copper.get('year', '未填写'),
                                    f"{copper.get('amount', 0)}%"
                                ])
                    
                    if copper_data:
                        self.add_table(copper_data, 
                                      ["年份", "金属铜回收率"], 
                                      "金属铜回收率")
                
                # 水资源重复利用率
                if 'waterReuseRate' in pollution_control:
                    reuse_data = []
                    reuse_info = pollution_control['waterReuseRate']
                    if isinstance(reuse_info, list):
                        for reuse in reuse_info:
                            if isinstance(reuse, dict):
                                reuse_data.append([
                                    reuse.get('year', '未填写'),
                                    f"{reuse.get('rate', 0)}%"
                                ])
                    
                    if reuse_data:
                        self.add_table(reuse_data, 
                                      ["年份", "水资源重复利用率"], 
                                      "水资源重复利用率")
        
        # 固体废物管理
        if pre_audit_data.solid_waste:
            self.add_title("3.6 工业固体废物管理", 2)
            
            solid_waste = pre_audit_data.solid_waste
            if isinstance(solid_waste, dict):
                # 一般固体废物
                if 'general' in solid_waste:
                    general_data = []
                    general_info = solid_waste['general']
                    if isinstance(general_info, list):
                        for waste in general_info:
                            if isinstance(waste, dict):
                                general_data.append([
                                    waste.get('year', '未填写'),
                                    waste.get('name', '未填写'),
                                    waste.get('type', '未填写'),
                                    f"{waste.get('amount', 0)}吨"
                                ])
                    
                    if general_data:
                        self.add_table(general_data, 
                                      ["年份", "废物名称", "废物类型", "产生量"], 
                                      "一般固体废物产生情况")
                
                # 危险废物
                if 'hazardous' in solid_waste:
                    hazardous_data = []
                    hazardous_info = solid_waste['hazardous']
                    if isinstance(hazardous_info, list):
                        for waste in hazardous_info:
                            if isinstance(waste, dict):
                                hazardous_data.append([
                                    waste.get('year', '未填写'),
                                    waste.get('name', '未填写'),
                                    waste.get('type', '未填写'),
                                    waste.get('code', '未填写'),
                                    f"{waste.get('amount', 0)}吨"
                                ])
                    
                    if hazardous_data:
                        self.add_table(hazardous_data, 
                                      ["年份", "废物名称", "废物类型", "废物代码", "产生量"], 
                                      "危险废物产生情况")
        
        # 自行监测情况
        if pre_audit_data.self_monitoring:
            self.add_title("3.7 自行监测情况", 2)
            
            self_monitoring = pre_audit_data.self_monitoring
            if isinstance(self_monitoring, dict):
                monitoring_data = []
                
                # 有组织废气监测
                if 'organizedGas' in self_monitoring:
                    gas_info = self_monitoring['organizedGas']
                    if isinstance(gas_info, dict):
                        monitoring_data.append([
                            "有组织废气",
                            gas_info.get('item', '未填写'),
                            f"{gas_info.get('concentration', 0)}mg/m³",
                            gas_info.get('point', '未填写'),
                            gas_info.get('standard', '未填写')
                        ])
                
                # 废水监测
                if 'wastewater' in self_monitoring:
                    water_info = self_monitoring['wastewater']
                    if isinstance(water_info, dict):
                        monitoring_data.append([
                            "废水",
                            water_info.get('item', '未填写'),
                            f"{water_info.get('concentration', 0)}mg/L",
                            water_info.get('point', '未填写'),
                            water_info.get('standard', '未填写')
                        ])
                
                # 噪声监测
                if 'noise' in self_monitoring:
                    noise_info = self_monitoring['noise']
                    if isinstance(noise_info, dict):
                        monitoring_data.append([
                            "噪声",
                            noise_info.get('item', '未填写'),
                            f"{noise_info.get('level', 0)}dB(A)",
                            noise_info.get('point', '未填写'),
                            noise_info.get('standard', '未填写')
                        ])
                
                if monitoring_data:
                    self.add_table(monitoring_data, 
                                  ["监测类型", "监测项目", "监测值", "监测点位", "执行标准"], 
                                  "自行监测情况")
    
    async def _generate_audit_results(self, audit_results: List, audit_report: Optional[PCBAuditReport]):
        """生成审核结果章节"""
        self.add_title("四、审核结果", 1)
        
        if audit_report:
            self.add_title("4.1 审核汇总", 2)
            
            summary_data = [
                ["总分", f"{audit_report.total_score or 0:.2f}"],
                ["综合等级", audit_report.overall_level or "未评估"],
                ["改进项数", f"{audit_report.improvement_items or 0}项"],
                ["限定性指标数", f"{audit_report.limiting_indicators_count or 0}项"],
                ["不达标限定性指标数", f"{audit_report.non_compliant_limiting_count or 0}项"],
                ["审核日期", audit_report.audit_date.strftime('%Y-%m-%d') if audit_report.audit_date else "未填写"],
                ["审核员", audit_report.auditor_name or "未填写"],
                ["报告状态", audit_report.status or "未填写"]
            ]
            
            self.add_table(summary_data, ["项目", "结果"], "审核汇总")
        
        if audit_results:
            self.add_title("4.2 详细审核结果", 2)
            
            # 按类别分组显示审核结果
            categories = {}
            for result in audit_results:
                # 获取指标信息
                indicator = await PCBIndicator.get_or_none(id=result.indicator_id)
                if indicator:
                    category = indicator.category
                    if category not in categories:
                        categories[category] = []
                    
                    categories[category].append({
                        'indicator_id': indicator.indicator_id,
                        'name': indicator.name,
                        'current_value': result.current_value,
                        'level': result.level,
                        'score': result.score,
                        'comment': result.comment
                    })
            
            # 显示各类别的审核结果
            for i, (category, results) in enumerate(categories.items()):
                self.add_title(f"4.2.{i + 1} {category}", 3)
                
                category_data = []
                for result in results:
                    category_data.append([
                        result['indicator_id'],
                        result['name'],
                        f"{result['current_value']:.2f}" if result['current_value'] is not None else "未填写",
                        result['level'] or "待评估",
                        f"{result['score']:.1f}" if result['score'] is not None else "未评分",
                        result['comment'] or "无"
                    ])
                
                self.add_table(category_data, 
                              ["指标编号", "指标名称", "当前值", "评级", "得分", "审核意见"], 
                              f"{category}审核结果")


# 创建报告生成器实例
report_generator = PCBReportGenerator()
