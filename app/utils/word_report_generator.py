#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PCB清洁生产审核报告生成器 - Word格式
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from typing import Dict, Any, List
import os


class PCBWordReportGenerator:
    """PCB清洁生产审核报告生成器"""
    
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()
    
    def setup_document_style(self):
        """设置文档样式"""
        # 设置默认字体为宋体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
    
    def add_title(self, title: str):
        """添加标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.runs[0]
        run.font.name = '黑体'
        run.font.size = Pt(22)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    def add_heading(self, text: str, level: int = 1):
        """添加小节标题"""
        heading = self.doc.add_heading(text, level=level)
        run = heading.runs[0]
        run.font.name = '黑体'
        run.font.size = Pt(16 if level == 1 else 14)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return heading
    
    def add_paragraph(self, text: str, indent: bool = False):
        """添加段落"""
        p = self.doc.add_paragraph(text)
        p.paragraph_format.line_spacing = 1.5
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        return p
    
    def add_table(self, headers: List[str], rows: List[List[Any]], col_widths: List[float] = None):
        """添加表格"""
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Light Grid Accent 1'
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 设置表头格式
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = '黑体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加数据行
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row_data):
                row_cells[col_idx].text = str(value) if value is not None else ""
                # 设置数据格式
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10.5)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 设置列宽
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)
        
        return table
    
    def generate_report(self, enterprise_data: Dict[str, Any], 
                       planning_data: Dict[str, Any],
                       preaudit_data: Dict[str, Any],
                       audit_data: Dict[str, Any]) -> str:
        """
        生成完整的审核报告
        
        Args:
            enterprise_data: 企业基本信息
            planning_data: 筹划与组织数据
            preaudit_data: 预审核数据
            audit_data: 审核数据
        
        Returns:
            生成的Word文档路径
        """
        
        # 1. 添加报告标题
        self.add_title("印制电路板行业清洁生产审核报告")
        self.doc.add_paragraph()
        
        # 2. 企业基本信息
        self._add_enterprise_info(enterprise_data)
        
        # 3. 筹划与组织
        self._add_planning_info(planning_data)
        
        # 4. 预审核数据
        self._add_preaudit_info(preaudit_data)
        
        # 5. 审核结果
        self._add_audit_results(audit_data)
        
        # 6. 改进建议
        self._add_improvement_suggestions(audit_data)
        
        # 7. 生成文件
        filename = f"PCB清洁生产审核报告_{enterprise_data['name']}_{datetime.now().strftime('%Y%m%d')}.docx"
        filepath = os.path.join('reports', filename)
        
        # 确保reports目录存在
        os.makedirs('reports', exist_ok=True)
        
        self.doc.save(filepath)
        return filepath
    
    def _add_enterprise_info(self, data: Dict[str, Any]):
        """添加企业基本信息部分"""
        self.add_heading("一、企业基本信息", level=1)
        
        # 辅助函数：安全转换为字符串
        def safe_str(value, default=''):
            """安全地将值转换为字符串"""
            if value is None:
                return default
            if isinstance(value, (int, float)):
                return str(value)
            return str(value) if value else default
        
        # 基本信息表格
        capital = data.get('capital')
        capacity = data.get('capacity')
        info_data = [
            ["企业名称", safe_str(data.get('name'))],
            ["统一社会信用代码", safe_str(data.get('unified_social_credit_code'))],
            ["所属地区", f"{safe_str(data.get('region'))} {safe_str(data.get('district'))}"],
            ["详细地址", safe_str(data.get('address'))],
            ["法人代表", safe_str(data.get('legal_representative'))],
            ["联系人", safe_str(data.get('contact_person'))],
            ["联系电话", safe_str(data.get('contact_phone'))],
            ["注册资本", f"{safe_str(capital)}万元" if capital is not None else ''],
            ["年产能", f"{safe_str(capacity)}万m²" if capacity is not None else ''],
            ["审核状态", safe_str(data.get('audit_status'))],
        ]
        
        self.add_table(["项目", "内容"], info_data, col_widths=[2.0, 4.5])
        self.doc.add_paragraph()
    
    def _add_planning_info(self, data: Dict[str, Any]):
        """添加筹划与组织信息"""
        self.add_heading("二、筹划与组织", level=1)
        
        # 2.1 审核领导小组
        self.add_heading("2.1 清洁生产审核领导小组", level=2)
        if data.get('leadership_team'):
            leadership_rows = []
            for member in data['leadership_team']:
                leadership_rows.append([
                    member.get('name', ''),
                    member.get('position', ''),
                    member.get('department', ''),
                    member.get('role', ''),
                    member.get('responsibilities', ''),
                    member.get('phone', '')
                ])
            self.add_table(
                ["姓名", "职位", "部门", "角色", "职责", "联系电话"],
                leadership_rows,
                col_widths=[0.8, 1.0, 1.0, 0.8, 2.0, 1.0]
            )
        self.doc.add_paragraph()
        
        # 2.2 工作小组
        self.add_heading("2.2 清洁生产工作小组", level=2)
        if data.get('work_team'):
            workteam_rows = []
            for member in data['work_team']:
                workteam_rows.append([
                    member.get('name', ''),
                    member.get('position', ''),
                    member.get('department', ''),
                    member.get('role', ''),
                    member.get('responsibilities', ''),
                    member.get('phone', '')
                ])
            self.add_table(
                ["姓名", "职位", "部门", "角色", "职责", "联系电话"],
                workteam_rows,
                col_widths=[0.8, 1.0, 1.0, 0.8, 2.0, 1.0]
            )
        self.doc.add_paragraph()
        
        # 2.3 工作计划
        self.add_heading("2.3 清洁生产审核工作计划", level=2)
        if data.get('work_plans'):
            workplan_rows = []
            for plan in data['work_plans']:
                workplan_rows.append([
                    f"第{plan.get('stage_order', '')}阶段",
                    plan.get('stage', ''),
                    plan.get('work_content', ''),
                    plan.get('planned_start_date', '')[:10] if plan.get('planned_start_date') else '',
                    plan.get('planned_end_date', '')[:10] if plan.get('planned_end_date') else '',
                    plan.get('responsible_department', ''),
                ])
            self.add_table(
                ["阶段", "阶段名称", "工作内容", "计划开始时间", "计划结束时间", "责任部门"],
                workplan_rows,
                col_widths=[0.8, 1.2, 2.0, 1.0, 1.0, 1.0]
            )
        self.doc.add_paragraph()
        
        # 2.4 培训记录
        self.add_heading("2.4 宣传与培训记录", level=2)
        if data.get('training_records'):
            training_rows = []
            for record in data['training_records']:
                training_rows.append([
                    record.get('title', ''),
                    record.get('date', '')[:10] if record.get('date') else '',
                    record.get('duration', ''),
                    record.get('participants', ''),
                    record.get('instructor', ''),
                    record.get('location', '')
                ])
            self.add_table(
                ["培训主题", "培训日期", "时长(分钟)", "参与人数", "讲师", "地点"],
                training_rows,
                col_widths=[2.0, 1.0, 0.8, 0.8, 0.8, 1.2]
            )
        self.doc.add_paragraph()
    
    def _add_preaudit_info(self, data: Dict[str, Any], include_tables: bool = True):
        """添加预审核数据（从独立表中收集的数据）"""
        self.add_heading("三、预审核数据", level=1)
        
        # 辅助函数：安全转换为数值字符串
        def safe_float_str(value, precision=2, default=''):
            """安全地将值转换为浮点数字符串"""
            if value is None or value == '':
                return default
            try:
                return f"{float(value):.{precision}f}"
            except (ValueError, TypeError):
                return default
        
        # 辅助函数：安全获取嵌套值
        def safe_get(data_dict, *keys, default=''):
            """安全地获取嵌套字典的值"""
            current = data_dict
            for key in keys:
                if isinstance(current, dict):
                    current = current.get(key)
                    if current is None:
                        return default
                else:
                    return default
            return current if current is not None else default
        
        # 3.1 企业总体生产情况
        self.add_heading("3.1 企业总体生产情况", level=2)
        
        # 产品产量表（近三年动态字段）
        production = data.get('production', {})
        if production.get('product_outputs') and include_tables:
            self.add_paragraph("（1）产品产量统计（近三年）")
            output_rows = []
            for item in production['product_outputs']:
                # 提取年份和产量数据
                year_data = []
                years_list = []
                outputs_list = []
                for year in ['2020', '2021', '2022', '2023', '2024']:
                    output_key = f'output_{year}'
                    val = item.get(output_key)
                    if val is not None and val != '':
                        years_list.append(year)
                        outputs_list.append(safe_float_str(val, 2))
                
                # 只显示有数据的年份
                if years_list:
                    years_str = '、'.join(years_list)
                    outputs_str = '、'.join(outputs_list)
                    output_rows.append([
                        str(item.get('type', '')),
                        str(item.get('main_product', '')),
                        str(item.get('layers', '')) if item.get('layers') else '',
                        str(item.get('unit', '') or 'm²'),
                        years_str,
                        outputs_str
                    ])
            
            if output_rows:
                self.add_table(
                    ["类型", "主要产品", "层数", "单位", "年份", "产量"],
                    output_rows,
                    col_widths=[1.0, 1.5, 0.8, 0.8, 1.5, 1.5]
                )
            self.doc.add_paragraph()
        
        # 合格率表（近三年动态字段）
        if production.get('qualification_rates') and include_tables:
            self.add_paragraph("（2）产品合格率（近三年）")
            qual_rows = []
            for item in production['qualification_rates']:
                rate_val = item.get('rate')
                rate_str = f"{safe_float_str(rate_val, 2)}%" if rate_val is not None else ''
                year_str = str(item.get('year', ''))
                if year_str and rate_str:
                    qual_rows.append([year_str, rate_str])
            
            if qual_rows:
                self.add_table(
                    ["年份", "合格率(%)"],
                    qual_rows,
                    col_widths=[2.0, 2.0]
                )
            self.doc.add_paragraph()
        
        # 产值表（近三年动态字段）
        if production.get('output_values') and include_tables:
            self.add_paragraph("（3）年产值情况（近三年）")
            value_rows = []
            for item in production['output_values']:
                year = str(item.get('year', ''))
                unit = str(item.get('unit', '') or '万元')
                annual_value = item.get('annual_output_value')
                income_tax = item.get('income_tax')
                
                if year:
                    value_rows.append([
                        year,
                        unit,
                        safe_float_str(annual_value, 2),
                        safe_float_str(income_tax, 2)
                    ])
            
            if value_rows:
                self.add_table(
                    ["年份", "单位", "年产值", "所得税"],
                    value_rows,
                    col_widths=[1.0, 1.0, 1.5, 1.5]
                )
            self.doc.add_paragraph()
        
        # 3.2 原辅材料使用情况（近三年动态字段）
        if data.get('raw_materials') and include_tables:
            self.add_heading("3.2 原辅材料使用情况", level=2)
            raw_material_rows = []
            for item in data['raw_materials']:
                # 提取年份数据
                year_data = []
                for year in ['2020', '2021', '2022', '2023', '2024']:
                    amount_key = f'amount_{year}'
                    unit_cons_key = f'unitConsumption_{year}'
                    val = item.get(amount_key)
                    if val is not None and val != '':
                        year_data.append({
                            'year': year,
                            'amount': safe_float_str(val, 2),
                            'unit_cons': safe_float_str(item.get(unit_cons_key), 4)
                        })
                
                # 构建行数据
                if year_data:
                    years_str = '、'.join([d['year'] for d in year_data])
                    amounts_str = '、'.join([d['amount'] for d in year_data])
                    unit_cons_str = '、'.join([d['unit_cons'] for d in year_data])
                    
                    raw_material_rows.append([
                        str(item.get('type', '')),
                        str(item.get('main_product', '')),
                        safe_float_str(item.get('product_output'), 2),
                        str(item.get('material_name', '')),
                        str(item.get('unit', '')),
                        years_str,
                        amounts_str,
                        unit_cons_str
                    ])
            
            if raw_material_rows:
                self.add_table(
                    ["类型", "主要产品", "产品产量(m²)", "原辅材料", "单位", "年份", "年消耗量", "单位产品消耗量"],
                    raw_material_rows,
                    col_widths=[0.8, 1.2, 1.0, 1.2, 0.6, 1.0, 1.2, 1.5]
                )
            self.doc.add_paragraph()
        
        # 3.3 主要工艺及装备使用
        if data.get('equipment') and include_tables:
            self.add_heading("3.3 主要工艺及装备使用", level=2)
            equipment_rows = []
            for item in data['equipment']:
                equipment_rows.append([
                    str(item.get('equipment_name', '')),
                    str(item.get('equipment_model', '')),
                    str(item.get('motor_model', '')),
                    safe_float_str(item.get('power'), 2) + 'kW' if item.get('power') is not None else '',
                    str(item.get('quantity', '')) if item.get('quantity') else '',
                    str(item.get('process', '')),
                    str(item.get('status', ''))
                ])
            if equipment_rows:
                self.add_table(
                    ["设备名称", "设备型号", "电机型号", "功率(kW)", "数量", "应用工艺", "运行状况"],
                    equipment_rows,
                    col_widths=[1.5, 1.5, 1.2, 0.8, 0.6, 1.2, 0.8]
                )
            self.doc.add_paragraph()
        
        # 3.4 资源能源消耗（近三年动态字段）
        resource_consumption = data.get('resource_consumption', {})
        if resource_consumption and include_tables:
            self.add_heading("3.4 资源能源消耗", level=2)
            
            # 用水情况
            if resource_consumption.get('water'):
                self.add_paragraph("（1）用水情况统计（近三年）")
                water_rows = []
                for item in resource_consumption['water']:
                    year_data = []
                    for year in ['2020', '2021', '2022', '2023', '2024']:
                        amount_key = f'amount_{year}'
                        val = item.get(amount_key)
                        if val is not None and val != '':
                            year_data.append(safe_float_str(val, 2))
                    
                    if year_data:
                        water_rows.append([
                            str(item.get('project', '')),
                            str(item.get('workshop', '') or ''),
                            str(item.get('unit', '')),
                            '、'.join(year_data)
                        ])
                
                if water_rows:
                    self.add_table(
                        ["项目", "使用车间", "单位", "年消耗量（按年份顺序）"],
                        water_rows,
                        col_widths=[1.0, 1.0, 0.8, 2.2]
                    )
                self.doc.add_paragraph()
            
            # 用电情况
            if resource_consumption.get('electricity'):
                self.add_paragraph("（2）用电情况统计（近三年）")
                electricity_rows = []
                for item in resource_consumption['electricity']:
                    year_data = []
                    for year in ['2020', '2021', '2022', '2023', '2024']:
                        amount_key = f'amount_{year}'
                        val = item.get(amount_key)
                        if val is not None and val != '':
                            year_data.append(safe_float_str(val, 2))
                    
                    if year_data:
                        electricity_rows.append([
                            str(item.get('project', '')),
                            str(item.get('workshop', '') or ''),
                            str(item.get('unit', '')),
                            '、'.join(year_data)
                        ])
                
                if electricity_rows:
                    self.add_table(
                        ["项目", "使用车间", "单位", "年消耗量（按年份顺序）"],
                        electricity_rows,
                        col_widths=[1.0, 1.0, 0.8, 2.2]
                    )
                self.doc.add_paragraph()
            
            # 天然气情况
            if resource_consumption.get('gas'):
                self.add_paragraph("（3）天然气使用情况（近三年）")
                gas_rows = []
                for item in resource_consumption['gas']:
                    year_data = []
                    for year in ['2020', '2021', '2022', '2023', '2024']:
                        amount_key = f'amount_{year}'
                        val = item.get(amount_key)
                        if val is not None and val != '':
                            year_data.append(safe_float_str(val, 2))
                    
                    if year_data:
                        gas_rows.append([
                            str(item.get('project', '')),
                            str(item.get('workshop', '') or ''),
                            str(item.get('unit', '')),
                            '、'.join(year_data)
                        ])
                
                if gas_rows:
                    self.add_table(
                        ["用途", "使用车间", "单位", "年消耗量（按年份顺序）"],
                        gas_rows,
                        col_widths=[1.0, 1.0, 0.8, 2.2]
                    )
                self.doc.add_paragraph()
        
        # 3.5 污染防治
        pollution_control = data.get('pollution_control', {})
        if pollution_control and include_tables:
            self.add_heading("3.5 污染防治", level=2)
            
            # 废水产生分析
            if pollution_control.get('wastewater_analysis'):
                self.add_paragraph("（1）废水产生分析")
                wastewater_rows = []
                for item in pollution_control['wastewater_analysis']:
                    wastewater_rows.append([
                        item.get('category', '') or '',
                        item.get('source', '') or '',
                        item.get('pollutants', '') or '',
                        item.get('disposal', '') or ''
                    ])
                if wastewater_rows:
                    self.add_table(
                        ["废水类别", "来源", "主要污染物", "处置方式"],
                        wastewater_rows,
                        col_widths=[1.2, 1.2, 1.5, 1.5]
                    )
                self.doc.add_paragraph()
            
            # 近三年废水情况统计
            if pollution_control.get('wastewater_stat'):
                self.add_paragraph("（2）近三年废水情况统计")
                wastewater_stat_rows = []
                for item in pollution_control['wastewater_stat']:
                    year_data = []
                    for year in ['2020', '2021', '2022', '2023', '2024']:
                        amount_key = f'amount_{year}'
                        val = item.get(amount_key)
                        if val is not None and val != '':
                            year_data.append(safe_float_str(val, 2))
                    
                    if year_data:
                        wastewater_stat_rows.append([
                            item.get('project', ''),
                            item.get('workshop', '') or '',
                            item.get('unit', ''),
                            '、'.join(year_data)
                        ])
                
                if wastewater_stat_rows:
                    self.add_table(
                        ["项目", "使用车间", "单位", "年产生量（按年份顺序）"],
                        wastewater_stat_rows,
                        col_widths=[1.0, 1.0, 0.8, 2.2]
                    )
                self.doc.add_paragraph()
            
            # 废气产生情况
            if pollution_control.get('waste_gas_analysis'):
                self.add_paragraph("（3）废气产生情况")
                wastegas_rows = []
                for item in pollution_control['waste_gas_analysis']:
                    wastegas_rows.append([
                        item.get('gas_type', '') or '',
                        item.get('pollutants', '') or '',
                        item.get('location', '') or '',
                        item.get('treatment', '') or ''
                    ])
                if wastegas_rows:
                    self.add_table(
                        ["废气种类", "主要污染物", "产生部位", "处理方法"],
                        wastegas_rows,
                        col_widths=[1.2, 1.5, 1.2, 1.5]
                    )
                self.doc.add_paragraph()
        
        # 3.6 工业固体废物管理（近三年动态字段）
        if data.get('solid_waste') and include_tables:
            self.add_heading("3.6 工业固体废物管理", level=2)
            waste_rows = []
            for item in data['solid_waste']:
                year_data = []
                for year in ['2020', '2021', '2022', '2023', '2024']:
                    amount_key = f'amount_{year}'
                    val = item.get(amount_key)
                    if val is not None and val != '':
                        year_data.append(safe_float_str(val, 2))
                
                if year_data:
                    waste_rows.append([
                        str(item.get('category', '')),
                        str(item.get('name', '')),
                        str(item.get('unit', '')),
                        '、'.join(year_data),
                        str(item.get('disposal_method', '') or '')
                    ])
            
            if waste_rows:
                self.add_table(
                    ["类别", "名称", "单位", "年产生量（按年份顺序）", "处置方式"],
                    waste_rows,
                    col_widths=[0.8, 1.2, 0.6, 2.0, 1.2]
                )
            self.doc.add_paragraph()
        
        # 3.7 自行监测情况
        self_monitoring = data.get('self_monitoring', {})
        if self_monitoring and include_tables:
            self.add_heading("3.7 自行监测情况", level=2)
            
            # 有组织废气监测
            if self_monitoring.get('organized_gas'):
                self.add_paragraph("（1）有组织废气监测情况")
                organized_gas_rows = []
                for item in self_monitoring['organized_gas']:
                    # 提取监测结果（动态字段）
                    results = []
                    monitoring_items = ['nitrogen_oxides', 'hydrogen_chloride', 'hydrogen_cyanide', 
                                       'sulfuric_acid_mist', 'chromic_acid_mist', 'fluoride',
                                       'phenol', 'non_methane_hydrocarbons', 'benzene', 'toluene',
                                       'xylene', 'toluene_xylene_total', 'vocs']
                    
                    for monitor_item in monitoring_items:
                        val = item.get(monitor_item)
                        if val is not None and val != '':
                            results.append(f"{monitor_item}: {safe_float_str(val, 2)}")
                    
                    organized_gas_rows.append([
                        str(item.get('monitoring_point', '') or ''),
                        str(item.get('monitoring_time', '') or ''),
                        '; '.join(results) if results else ''
                    ])
                
                if organized_gas_rows:
                    self.add_table(
                        ["监测点位", "监测时间", "监测结果(mg/m³)"],
                        organized_gas_rows,
                        col_widths=[1.0, 1.2, 3.2]
                    )
                self.doc.add_paragraph()
            
            # 无组织废气监测
            if self_monitoring.get('unorganized_gas'):
                self.add_paragraph("（2）无组织废气监测情况")
                unorganized_gas_rows = []
                for item in self_monitoring['unorganized_gas']:
                    unorganized_gas_rows.append([
                        str(item.get('sampling_point', '') or ''),
                        str(item.get('sampling_time', '') or ''),
                        str(item.get('monitoring_factor', '') or ''),
                        safe_float_str(item.get('emission_concentration'), 2),
                        safe_float_str(item.get('emission_limit'), 2),
                        str(item.get('compliance', '') or '')
                    ])
                
                if unorganized_gas_rows:
                    self.add_table(
                        ["采样点位", "采样时间", "监测因子", "排放浓度(mg/m³)", "排放限值(mg/m³)", "达标情况"],
                        unorganized_gas_rows,
                        col_widths=[1.0, 1.0, 1.0, 1.0, 1.0, 0.8]
                    )
                self.doc.add_paragraph()
            
            # 废水排放监测
            if self_monitoring.get('wastewater'):
                self.add_paragraph("（3）废水排放监测情况")
                wastewater_rows = []
                for item in self_monitoring['wastewater']:
                    # 提取监测结果（动态字段）
                    results = []
                    monitoring_items = ['ph', 'cod', 'ammonia_nitrogen', 'total_nitrogen', 'total_phosphorus',
                                       'copper', 'nickel', 'total_cyanide', 'nickel_outlet']
                    
                    for monitor_item in monitoring_items:
                        val = item.get(monitor_item)
                        if val is not None:
                            results.append(f"{monitor_item}: {safe_float_str(val, 2)}")
                    
                    wastewater_rows.append([
                        item.get('monitoring_point', '') or '',
                        item.get('monitoring_time', '') or '',
                        '; '.join(results) if results else ''
                    ])
                
                if wastewater_rows:
                    self.add_table(
                        ["监测点位", "监测时间", "监测结果(mg/L)"],
                        wastewater_rows,
                        col_widths=[1.0, 1.2, 3.2]
                    )
                self.doc.add_paragraph()
            
            # 废气排放监测
            if self_monitoring.get('gas_emission'):
                self.add_paragraph("（4）废气排放监测情况")
                gas_emission_rows = []
                for item in self_monitoring['gas_emission']:
                    gas_emission_rows.append([
                        item.get('detection_point', '') or '',
                        item.get('detection_item', '') or '',
                        safe_float_str(item.get('detection_result'), 2),
                        safe_float_str(item.get('permitted_emission_limit'), 2),
                        safe_float_str(item.get('stack_height'), 1)
                    ])
                
                if gas_emission_rows:
                    self.add_table(
                        ["检测点位", "检测项目", "检测结果", "许可排放浓度限值", "排气筒高(m)"],
                        gas_emission_rows,
                        col_widths=[1.0, 1.2, 1.0, 1.2, 0.8]
                    )
                self.doc.add_paragraph()
            
            # 噪声监测
            if self_monitoring.get('noise'):
                self.add_paragraph("（5）近三年厂界噪声监测情况")
                noise_rows = []
                for item in self_monitoring['noise']:
                    noise_rows.append([
                        str(item.get('monitoring_point', '') or ''),
                        str(item.get('monitoring_time', '') or ''),
                        safe_float_str(item.get('daytime_result'), 1),
                        safe_float_str(item.get('daytime_standard'), 1),
                        safe_float_str(item.get('nighttime_result'), 1),
                        safe_float_str(item.get('nighttime_standard'), 1)
                    ])
                
                if noise_rows:
                    self.add_table(
                        ["监测点位", "监测时间", "昼间结果dB(A)", "昼间标准dB(A)", "夜间结果dB(A)", "夜间标准dB(A)"],
                        noise_rows,
                        col_widths=[1.0, 1.2, 1.0, 1.0, 1.0, 1.0]
                    )
                self.doc.add_paragraph()
    
    def _add_audit_results(self, data: Dict[str, Any], include_tables: bool = True):
        """添加审核结果"""
        self.add_heading("四、审核结果", level=1)
        
        # 辅助函数：安全转换为数值字符串
        def safe_float_str(value, precision=2, default=''):
            """安全地将值转换为浮点数字符串"""
            if value is None or value == '':
                return default
            try:
                return f"{float(value):.{precision}f}"
            except (ValueError, TypeError):
                return default
        
        # 审核汇总
        if data.get('summary'):
            summary = data['summary']
            total_score = summary.get('total_score')
            summary_text = f"综合得分：{safe_float_str(total_score, 2)}分\n" if total_score is not None else ""
            summary_text += f"综合等级：{summary.get('overall_level', '')}\n" if summary.get('overall_level') else ""
            summary_text += f"待改进项：{summary.get('improvement_items', 0)}项\n" if summary.get('improvement_items') is not None else ""
            summary_text += f"限定性指标不达标：{summary.get('non_compliant_limiting_count', 0)}项" if summary.get('non_compliant_limiting_count') is not None else ""
            
            if summary_text:
                self.add_paragraph(summary_text.strip())
            self.doc.add_paragraph()
        
        # 详细审核结果表
        if data.get('results') and include_tables:
            self.add_heading("4.1 详细审核结果", level=2)
            
            # 按类别分组显示
            categories = {}
            for result in data['results']:
                indicator = result.get('indicator', {})
                category = indicator.get('category', '未分类') if indicator else '未分类'
                if category not in categories:
                    categories[category] = []
                categories[category].append(result)
            
            for category, results in categories.items():
                self.add_heading(category, level=3)
                indicator_rows = []
                for item in results:
                    ind = item.get('indicator', {})
                    current_value = item.get('current_value')
                    level = item.get('level', '')
                    score = item.get('score')
                    
                    indicator_rows.append([
                        str(ind.get('indicator_id', '')) if ind else '',
                        ind.get('name', '') if ind else '',
                        f"{float(current_value):.4f}" if current_value is not None else '',
                        level,
                        f"{float(score):.2f}" if score is not None else '0.00'
                    ])
                
                if indicator_rows:
                    self.add_table(
                        ["指标编号", "指标名称", "当前值", "评级", "得分"],
                        indicator_rows,
                        col_widths=[0.8, 3.0, 1.0, 0.8, 0.8]
                    )
                self.doc.add_paragraph()
    
    def _add_improvement_suggestions(self, data: Dict[str, Any]):
        """添加改进建议"""
        self.add_heading("六、清洁生产改进建议", level=1)
        
        # 收集所有不达标或低等级的指标
        needs_improvement = []
        if data.get('results'):
            for item in data['results']:
                level = item.get('level', '')
                if level in ['II级', 'III级', '不达标']:
                    needs_improvement.append(item)
        
        if needs_improvement:
            self.add_paragraph(f"根据审核结果，共有{len(needs_improvement)}项指标需要改进。针对这些指标，提出以下改进建议：")
            self.doc.add_paragraph()
            
            for idx, item in enumerate(needs_improvement, 1):
                ind = item.get('indicator', {})
                
                self.add_paragraph(f"{idx}. {ind.get('name', '') if ind else '未知指标'}（当前评级：{item.get('level', '')}）")
                
                comment = item.get('comment', '')
                if comment:
                    self.add_paragraph(f"   审核意见：{comment}", indent=True)
                
                self.doc.add_paragraph()
        else:
            self.add_paragraph("恭喜！所有指标均达到I级标准，无需改进。")
        
        self.doc.add_paragraph()
        
        # 添加报告结尾
        self.add_heading("七、结论", level=1)
        self.add_paragraph("本次清洁生产审核全面评估了企业的生产工艺、资源消耗、污染排放等各方面情况。", indent=True)
        self.add_paragraph("企业应根据本报告提出的改进建议，制定具体的实施计划，持续推进清洁生产工作。", indent=True)
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # 添加签章区域
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"审核日期：{datetime.now().strftime('%Y年%m月%d日')}")
    
    def get_current_time(self):
        """获取当前时间"""
        return datetime.now()
    
    async def generate_complete_report(
        self,
        enterprise_id: int,
        enterprise_data: Dict[str, Any] = None,
        planning_data: Dict[str, Any] = None,
        preaudit_data: Dict[str, Any] = None,
        audit_data: Dict[str, Any] = None,
        problem_solution_data: Dict[str, Any] = None,
        include_tables: bool = True,
        include_recommendations: bool = True
    ) -> str:
        """
        生成完整的审核报告（包含所有模块）
        
        Args:
            enterprise_id: 企业ID
            enterprise_data: 企业基本信息
            planning_data: 筹划与组织数据
            preaudit_data: 预审核数据
            audit_data: 审核数据
            problem_solution_data: 问题及清洁生产方案数据
            include_tables: 是否包含详细表格
            include_recommendations: 是否包含改进建议
        
        Returns:
            生成的Word文档路径
        """
        # 重新初始化文档
        self.doc = Document()
        self.setup_document_style()
        
        # 1. 添加报告标题
        self.add_title("印制电路板行业清洁生产审核评估报告")
        self.doc.add_paragraph()
        
        # 2. 企业基本信息
        if enterprise_data:
            self._add_enterprise_info(enterprise_data)
        
        # 3. 筹划与组织
        if planning_data:
            self._add_planning_info(planning_data)
        
        # 4. 预审核数据
        if preaudit_data:
            self._add_preaudit_info(preaudit_data, include_tables)
        
        # 5. 审核结果
        if audit_data:
            self._add_audit_results(audit_data, include_tables)
        
        # 6. 问题及清洁生产方案（新增）
        if problem_solution_data:
            self._add_problem_solution_info(problem_solution_data, include_tables)
        
        # 7. 改进建议
        if include_recommendations and audit_data:
            self._add_improvement_suggestions(audit_data)
        
        # 8. 生成文件
        enterprise_name = enterprise_data.get('name', 'Unknown') if enterprise_data else f'Enterprise_{enterprise_id}'
        filename = f"PCB审核报告_{enterprise_name}_{enterprise_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join('reports', filename)
        
        # 确保reports目录存在
        os.makedirs('reports', exist_ok=True)
        
        self.doc.save(filepath)
        return filepath
    
    def _add_problem_solution_info(self, data: Dict[str, Any], include_tables: bool = True):
        """添加问题及清洁生产方案部分"""
        self.add_heading("五、问题及清洁生产方案", level=1)
        
        # 5.1 问题清单
        self.add_heading("5.1 问题清单（Ⅱ级及以下指标）", level=2)
        if data.get('issues') and include_tables:
            issues = data['issues']
            issue_rows = []
            for idx, issue in enumerate(issues, 1):
                issue_rows.append([
                    str(idx),
                    issue.get('primary_indicator', ''),
                    str(issue.get('primary_weight', '')) if issue.get('primary_weight') else '',
                    issue.get('secondary_indicator', ''),
                    str(issue.get('secondary_weight', '')) if issue.get('secondary_weight') else '',
                    issue.get('current_level', ''),
                    issue.get('problem', '') or '',
                    issue.get('advice', '') or '',
                    issue.get('department', '') or '',
                    issue.get('owner', '') or '',
                    issue.get('deadline', '') or ''
                ])
            self.add_table(
                ["序号", "一级指标", "一级权重", "二级指标", "二级权重", "当前评级", "问题描述", "整改建议", "责任部门", "责任人", "整改期限"],
                issue_rows,
                col_widths=[0.5, 1.2, 0.8, 1.5, 0.8, 0.8, 1.5, 1.5, 0.8, 0.8, 1.0]
            )
        self.doc.add_paragraph()
        
        # 5.2 权重总和计分排序
        if data.get('scoring'):
            self.add_heading("5.2 权重总和计分排序", level=2)
            scoring = data['scoring']
            
            if scoring.get('factors') and include_tables:
                factors = scoring['factors']
                factor_rows = [[f.get('name', ''), str(f.get('weight', ''))] for f in factors]
                self.add_table(["因素", "权重W"], factor_rows, col_widths=[3.0, 1.0])
                self.doc.add_paragraph()
            
            if scoring.get('rankings') and include_tables:
                rankings = scoring['rankings']
                ranking_rows = [[str(idx + 1), r.get('name', ''), str(r.get('total_score', ''))] 
                               for idx, r in enumerate(rankings)]
                self.add_table(["排名", "审核重点", "总分"], ranking_rows, col_widths=[0.8, 3.0, 1.2])
            self.doc.add_paragraph()
        
        # 5.3 无/低费方案库
        self.add_heading("5.3 无/低费方案库", level=2)
        if data.get('low_cost_schemes') and include_tables:
            schemes = data['low_cost_schemes']
            scheme_rows = []
            for idx, scheme in enumerate(schemes, 1):
                scheme_rows.append([
                    str(idx),
                    scheme.get('name', ''),
                    scheme.get('intro', '') or '',
                    scheme.get('economic_benefit', '') or '',
                    scheme.get('environment_benefit', '') or ''
                ])
            self.add_table(
                ["序号", "方案名称", "方案简介", "经济效益", "环境效益"],
                scheme_rows,
                col_widths=[0.5, 1.5, 2.0, 1.5, 1.5]
            )
        self.doc.add_paragraph()
        
        # 5.4 中/高费方案库
        self.add_heading("5.4 中/高费方案库", level=2)
        if data.get('medium_high_cost_schemes') and include_tables:
            schemes = data['medium_high_cost_schemes']
            scheme_rows = []
            for idx, scheme in enumerate(schemes, 1):
                scheme_rows.append([
                    str(idx),
                    scheme.get('name', ''),
                    scheme.get('intro', '') or '',
                    scheme.get('cost_level', ''),
                    scheme.get('economic_benefit', '') or '',
                    scheme.get('environment_benefit', '') or ''
                ])
            self.add_table(
                ["序号", "方案名称", "方案简介", "费用等级", "经济效益", "环境效益"],
                scheme_rows,
                col_widths=[0.5, 1.5, 2.0, 0.8, 1.3, 1.3]
            )
        self.doc.add_paragraph()


# 创建全局实例
generator = PCBWordReportGenerator()


# 导出的工具函数
async def generate_pcb_word_report(
    enterprise_data: Dict[str, Any],
    planning_data: Dict[str, Any],
    preaudit_data: Dict[str, Any],
    audit_data: Dict[str, Any]
) -> str:
    """
    生成PCB清洁生产审核Word报告
    
    Returns:
        生成的Word文档路径
    """
    filepath = generator.generate_report(
        enterprise_data=enterprise_data,
        planning_data=planning_data,
        preaudit_data=preaudit_data,
        audit_data=audit_data
    )
    return filepath

